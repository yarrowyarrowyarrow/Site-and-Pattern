"""
export_plant_docx.py — Export the full plant catalogue to plant_data.docx.

Run from project root:
    python scripts/export_plant_docx.py

Produces plant_data.docx with three sections:
  1. Data Field Reference  — every field name, label, and valid values
  2. Plant Catalogue       — all plants, one row per plant, all core fields
  3. Planting Calendar     — all plants with Jan–Dec monthly status
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def _set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_header_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def _set_col_widths(table, widths_cm: list[float]):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _small_font(table, size_pt: float = 7):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


# ── field reference data ──────────────────────────────────────────────────────

FIELD_REFERENCE = [
    ("common_name",        "Common Name",            "Text — English common name"),
    ("scientific_name",    "Scientific Name",         "Text — Latin binomial"),
    ("plant_type",         "Plant Type",              "tree | shrub | herb | groundcover | vine"),
    ("hardiness_zone_min", "Hardiness Zone Min",      "Integer 1–9 (USDA/Canadian zone)"),
    ("hardiness_zone_max", "Hardiness Zone Max",      "Integer 1–9"),
    ("sun_requirement",    "Sun Requirement",         "full_sun | partial_shade | full_shade"),
    ("water_needs",        "Water Needs",             "low | medium | high"),
    ("soil_ph_min",        "Soil pH Min",             "Real number (e.g. 5.0)"),
    ("soil_ph_max",        "Soil pH Max",             "Real number (e.g. 7.5)"),
    ("mature_height_m",    "Mature Height (m)",       "Real number — height at full maturity"),
    ("spacing_m",          "Spacing (m)",             "Real number — recommended planting spacing"),
    ("deciduous_evergreen","Foliage Retention",       "deciduous | evergreen | herbaceous"),
    ("perennial_annual",   "Life Cycle",              "perennial | annual | biennial"),
    ("growth_rate",        "Growth Rate",             "slow | moderate | fast"),
    ("years_to_maturity",  "Years to Maturity",       "Integer — years to reach full size"),
    ("growth_curve",       "Growth Curve",            "fast_early | steady | slow_start"),
    ("bloom_period",       "Bloom Period",            "Text (e.g. \"May–June\")"),
    ("fruit_period",       "Fruit Period",            "Text (e.g. \"August–September\")"),
    ("permaculture_uses",  "Permaculture Uses",       "Comma-separated: biomass, dynamic_accumulator, food_forest, groundcover, medicinal, nitrogen_fixer, pest_repellent, pioneer, pollinator, wildlife_habitat, windbreak"),
    ("native_region",      "Native Region",           "Text (e.g. \"Western Canada\")"),
    ("native_to_alberta",  "Native to Alberta",       "1 = yes, 0 = no"),
    ("edible_parts",       "Edible Parts",            "Comma-separated: bark, buds, bulb, corms, flowers, fruit, leaves, needles (tea), nuts, pollen, resin, root (tea), roots, sap, seeds, shoots, shoots (cooked), tuber"),
    ("notes",              "Notes",                   "Free text — extended description / warnings"),
    ("marker_color",       "Marker Color",            "Hex colour string (optional, e.g. #4CAF50) — map marker override"),
    ("cal_jan",            "January",                 "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_feb",            "February",                "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_mar",            "March",                   "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_apr",            "April",                   "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_may",            "May",                     "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_jun",            "June",                    "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_jul",            "July",                    "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_aug",            "August",                  "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_sep",            "September",               "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_oct",            "October",                 "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_nov",            "November",                "dormant | direct_sow | growing | harvest | pruning"),
    ("cal_dec",            "December",                "dormant | direct_sow | growing | harvest | pruning"),
]

HEADER_BG = "1b5e20"   # dark green
HEADER_FG = RGBColor(0xC8, 0xE6, 0xC9)  # light green
ALT_ROW_BG = "f1f8e9"  # very light green


# ── section builders ──────────────────────────────────────────────────────────

def _build_field_reference(doc: Document):
    _add_heading(doc, "1. Data Field Reference")
    doc.add_paragraph(
        "Every field stored per plant, its human-readable label, and the set of valid values."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0]
    _shade_cell(hdr.cells[0], HEADER_BG)
    _shade_cell(hdr.cells[1], HEADER_BG)
    _shade_cell(hdr.cells[2], HEADER_BG)
    for cell, text in zip(hdr.cells, ["Field (JSON key)", "Human Label", "Valid Values / Description"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = HEADER_FG
        run.font.size = Pt(9)

    for i, (field, label, values) in enumerate(FIELD_REFERENCE):
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(field).font.size = Pt(8)
        row.cells[1].paragraphs[0].add_run(label).font.size = Pt(8)
        row.cells[2].paragraphs[0].add_run(values).font.size = Pt(8)
        if i % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, ALT_ROW_BG)

    col_widths = [5.5, 4.5, 14.0]
    _set_col_widths(table, col_widths)


def _build_plant_catalogue(doc: Document, plants: list[dict]):
    doc.add_page_break()
    _add_heading(doc, "2. Plant Catalogue")
    doc.add_paragraph(f"All {len(plants)} plants — core fields. Calendar data is in Section 3.")

    headers = [
        "Common Name", "Scientific Name", "Type", "Zone\n(min–max)",
        "Sun", "Water", "Ht\n(m)", "Spc\n(m)", "pH Range",
        "Foliage", "Life\nCycle", "Growth\nRate", "Yrs\nMature",
        "Native\nAB", "Edible Parts", "Permaculture Uses", "Notes",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for cell, text in zip(hdr_row.cells, headers):
        _shade_cell(cell, HEADER_BG)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = HEADER_FG
        run.font.size = Pt(7)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sorted_plants = sorted(plants, key=lambda p: p.get("common_name", ""))

    for i, p in enumerate(sorted_plants):
        zone_min = p.get("hardiness_zone_min", "")
        zone_max = p.get("hardiness_zone_max", "")
        zone = f"{zone_min}–{zone_max}" if zone_min and zone_max else zone_min or zone_max or ""

        ph_min = p.get("soil_ph_min", "")
        ph_max = p.get("soil_ph_max", "")
        ph = f"{ph_min}–{ph_max}" if ph_min and ph_max else ph_min or ph_max or ""

        native_ab = "Yes" if p.get("native_to_alberta") == 1 else "No"

        values = [
            p.get("common_name", ""),
            p.get("scientific_name", ""),
            p.get("plant_type", "").title(),
            zone,
            p.get("sun_requirement", "").replace("_", " ").title(),
            p.get("water_needs", "").title(),
            str(p.get("mature_height_m", "")),
            str(p.get("spacing_m", "")),
            ph,
            p.get("deciduous_evergreen", "").title(),
            p.get("perennial_annual", "").title(),
            p.get("growth_rate", "").title(),
            str(p.get("years_to_maturity", "")),
            native_ab,
            p.get("edible_parts", ""),
            p.get("permaculture_uses", ""),
            p.get("notes", ""),
        ]

        row = table.add_row()
        if i % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, ALT_ROW_BG)
        for cell, val in zip(row.cells, values):
            cell.paragraphs[0].add_run(val).font.size = Pt(7)

    # Column widths (cm) — tuned for landscape Letter (27.94 cm usable)
    col_widths = [3.2, 3.8, 1.6, 1.4, 2.0, 1.4, 0.9, 0.9, 1.4, 1.8, 1.4, 1.4, 1.0, 1.0, 2.5, 3.2, 4.1]
    _set_col_widths(table, col_widths)


def _build_calendar(doc: Document, plants: list[dict]):
    doc.add_page_break()
    _add_heading(doc, "3. Planting Calendar")
    doc.add_paragraph(
        "Monthly status for each plant. "
        "D = dormant  |  G = growing  |  H = harvest  |  S = direct_sow  |  P = pruning"
    )

    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    cal_keys = [f"cal_{m.lower()}" for m in months]

    STATUS_ABBREV = {
        "dormant":     "D",
        "growing":     "G",
        "harvest":     "H",
        "direct_sow":  "S",
        "pruning":     "P",
        "":            "",
    }

    STATUS_COLOR = {
        "D": "eeeeee",
        "G": "c8e6c9",
        "H": "fff9c4",
        "S": "bbdefb",
        "P": "ffe0b2",
        "":  "ffffff",
    }

    headers = ["Common Name"] + months
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for cell, text in zip(hdr_row.cells, headers):
        _shade_cell(cell, HEADER_BG)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = HEADER_FG
        run.font.size = Pt(7)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sorted_plants = sorted(plants, key=lambda p: p.get("common_name", ""))

    for p in sorted_plants:
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(p.get("common_name", "")).font.size = Pt(7)
        for j, key in enumerate(cal_keys):
            status = p.get(key, "")
            abbrev = STATUS_ABBREV.get(status, status[:1].upper() if status else "")
            cell = row.cells[j + 1]
            _shade_cell(cell, STATUS_COLOR.get(abbrev, "ffffff"))
            run = cell.paragraphs[0].add_run(abbrev)
            run.font.size = Pt(7)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # name col wide, month cols narrow
    col_widths = [5.5] + [1.8] * 12
    _set_col_widths(table, col_widths)


# ── main ──────────────────────────────────────────────────────────────────────

def main(out_path: str | None = None) -> str:
    """Build the plant-catalogue DOCX and return the path written.

    Args:
        out_path: where to write the .docx. Defaults to
            ``<project_root>/plant_data.docx`` (the original CLI
            behaviour). The headless scripting API
            (src.permadesign_api.export_plant_catalogue_docx) passes an
            explicit path here.
    """
    project_root = Path(__file__).resolve().parent.parent
    plants_path = project_root / "data" / "plants_master.json"
    output_path = Path(out_path) if out_path else (project_root / "plant_data.docx")

    with open(plants_path, encoding="utf-8") as f:
        plants = json.load(f)

    doc = Document()

    # Landscape for all sections
    section = doc.sections[0]
    _set_landscape(section)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title = doc.add_heading("PermaDesign — Plant Data Export", level=0)
    title.paragraph_format.space_after = Pt(4)
    subtitle = doc.add_paragraph(
        f"Source: data/plants.json  |  {len(plants)} plants  |  Generated 2026-04-23"
    )
    subtitle.paragraph_format.space_after = Pt(12)

    _build_field_reference(doc, )
    _build_plant_catalogue(doc, plants)
    _build_calendar(doc, plants)

    doc.save(str(output_path))
    print(f"Saved: {output_path}  ({len(plants)} plants)")
    return str(output_path)


if __name__ == "__main__":
    main()
